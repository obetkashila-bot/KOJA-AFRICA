import os
import io
import uuid
import logging
from datetime import datetime, timezone
from functools import wraps
from html import escape

import requests
from dotenv import load_dotenv

from flask import (
    Flask,
    request,
    redirect,
    url_for,
    session,
    render_template_string,
    flash,
    send_file,
    abort,
    jsonify,
)

from werkzeug.security import (
    generate_password_hash,
    check_password_hash,
)

from werkzeug.utils import secure_filename

from reportlab.lib.pagesizes import A4
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
)
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER

from docx import Document
from openpyxl import Workbook


# ============================================================
# LOAD ENVIRONMENT
# ============================================================

load_dotenv()


# ============================================================
# FLASK APP
# ============================================================

app = Flask(__name__)

app.secret_key = os.environ.get(
    "SECRET_KEY",
    "CHANGE_THIS_SECRET_KEY_IN_RENDER",
)

app.config["MAX_CONTENT_LENGTH"] = 15 * 1024 * 1024


# ============================================================
# LOGGING
# ============================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
)


# ============================================================
# SUPABASE CONFIGURATION
# ============================================================

SUPABASE_URL = os.environ.get(
    "SUPABASE_URL",
    "",
).strip().rstrip("/")


SUPABASE_SERVICE_KEY = os.environ.get(
    "SUPABASE_SERVICE_KEY",
    "",
).strip()


# ============================================================
# APPLICATION CONFIGURATION
# ============================================================

APP_NAME = "KOJA AFRICA"

APP_TAGLINE = (
    "Assignment Questions • Academic Answers • "
    "Learning Resources"
)


STORAGE_BUCKET = os.environ.get(
    "SUPABASE_STORAGE_BUCKET",
    "koja-assignments",
).strip()


# ============================================================
# ADMIN LOGIN
# ============================================================

ADMIN_EMAIL = os.environ.get(
    "ADMIN_EMAIL",
    "admin@koja-africa.com",
).strip().lower()


ADMIN_PASSWORD = os.environ.get(
    "ADMIN_PASSWORD",
    "ChangeThisAdminPassword123!",
)


# ============================================================
# FILE SETTINGS
# ============================================================

ALLOWED_EXTENSIONS = {
    "pdf",
    "doc",
    "docx",
    "xls",
    "xlsx",
    "png",
    "jpg",
    "jpeg",
    "webp",
    "txt",
}

MAX_FILE_SIZE = 15 * 1024 * 1024


# ============================================================
# BASIC HELPERS
# ============================================================

def utc_now():
    return datetime.now(timezone.utc).isoformat()


def clean(value):
    """
    Safely escape text before putting it into HTML.
    """

    if value is None:
        return ""

    return escape(str(value))


def allowed_file(filename):
    if not filename:
        return False

    if "." not in filename:
        return False

    extension = filename.rsplit(
        ".",
        1,
    )[1].lower()

    return extension in ALLOWED_EXTENSIONS


# ============================================================
# SUPABASE HEADERS
# ============================================================

def supabase_headers():
    if not SUPABASE_SERVICE_KEY:
        raise RuntimeError(
            "SUPABASE_SERVICE_KEY is missing."
        )

    return {
        "apikey": SUPABASE_SERVICE_KEY,
        "Authorization": (
            f"Bearer {SUPABASE_SERVICE_KEY}"
        ),
        "Content-Type": "application/json",
    }


# ============================================================
# SUPABASE REST REQUEST
# ============================================================

def supabase_request(
    method,
    table,
    params=None,
    data=None,
    headers=None,
    timeout=30,
):
    """
    Generic Supabase REST request.

    Uses the Supabase service key server-side.
    """

    if not SUPABASE_URL:
        raise RuntimeError(
            "SUPABASE_URL is missing in Render environment variables."
        )

    if not SUPABASE_SERVICE_KEY:
        raise RuntimeError(
            "SUPABASE_SERVICE_KEY is missing in Render environment variables."
        )

    url = (
        f"{SUPABASE_URL}/rest/v1/{table}"
    )

    final_headers = supabase_headers()

    if headers:
        final_headers.update(headers)

    try:

        response = requests.request(
            method=method,
            url=url,
            params=params,
            json=data,
            headers=final_headers,
            timeout=timeout,
        )

    except requests.RequestException as exc:

        logging.exception(
            "Supabase network error."
        )

        raise RuntimeError(
            f"Could not connect to Supabase: {exc}"
        )

    if not response.ok:

        logging.error(
            "Supabase request failed. "
            "Method=%s Table=%s Status=%s Response=%s",
            method,
            table,
            response.status_code,
            response.text,
        )

        raise RuntimeError(
            f"Database request failed: "
            f"{response.status_code} "
            f"{response.text}"
        )

    if not response.text:
        return []

    try:
        return response.json()

    except ValueError:
        return response.text


# ============================================================
# DATABASE SELECT
# ============================================================

def db_select(
    table,
    columns="*",
    filters=None,
    limit=None,
    order=None,
):
    params = {
        "select": columns,
    }

    if filters:
        params.update(filters)

    if limit is not None:
        params["limit"] = str(limit)

    if order:
        params["order"] = order

    return supabase_request(
        "GET",
        table,
        params=params,
    )


# ============================================================
# DATABASE INSERT
# ============================================================

def db_insert(
    table,
    data,
    returning=True,
):
    if returning:

        headers = {
            "Prefer": "return=representation",
        }

    else:

        headers = {
            "Prefer": "return=minimal",
        }

    return supabase_request(
        "POST",
        table,
        data=data,
        headers=headers,
    )


# ============================================================
# DATABASE UPDATE
# ============================================================

def db_update(
    table,
    filters,
    data,
):
    return supabase_request(
        "PATCH",
        table,
        params=filters,
        data=data,
        headers={
            "Prefer": "return=representation",
        },
    )


# ============================================================
# DATABASE DELETE
# ============================================================

def db_delete(
    table,
    filters,
):
    return supabase_request(
        "DELETE",
        table,
        params=filters,
        headers={
            "Prefer": "return=minimal",
        },
    )


# ============================================================
# STORAGE UPLOAD
# ============================================================

def storage_upload(
    file_bytes,
    storage_path,
    content_type,
):
    if not SUPABASE_URL:
        raise RuntimeError(
            "SUPABASE_URL is missing."
        )

    if not SUPABASE_SERVICE_KEY:
        raise RuntimeError(
            "SUPABASE_SERVICE_KEY is missing."
        )

    if not STORAGE_BUCKET:
        raise RuntimeError(
            "SUPABASE_STORAGE_BUCKET is missing."
        )

    url = (
        f"{SUPABASE_URL}/storage/v1/object/"
        f"{STORAGE_BUCKET}/{storage_path}"
    )

    headers = {
        "Authorization": (
            f"Bearer {SUPABASE_SERVICE_KEY}"
        ),
        "apikey": SUPABASE_SERVICE_KEY,
        "Content-Type": (
            content_type
            or "application/octet-stream"
        ),
        "x-upsert": "true",
    }

    try:

        response = requests.post(
            url,
            headers=headers,
            data=file_bytes,
            timeout=60,
        )

    except requests.RequestException as exc:

        raise RuntimeError(
            f"Storage connection failed: {exc}"
        )

    if not response.ok:

        raise RuntimeError(
            "Storage upload failed: "
            f"{response.status_code} "
            f"{response.text}"
        )

    return storage_path


# ============================================================
# STORAGE DOWNLOAD
# ============================================================

def storage_download(storage_path):

    if not storage_path:
        raise RuntimeError(
            "Storage path is missing."
        )

    url = (
        f"{SUPABASE_URL}/storage/v1/object/"
        f"{STORAGE_BUCKET}/{storage_path}"
    )

    headers = {
        "Authorization": (
            f"Bearer {SUPABASE_SERVICE_KEY}"
        ),
        "apikey": SUPABASE_SERVICE_KEY,
    }

    try:

        response = requests.get(
            url,
            headers=headers,
            timeout=60,
        )

    except requests.RequestException as exc:

        raise RuntimeError(
            f"Storage connection failed: {exc}"
        )

    if not response.ok:

        raise RuntimeError(
            "Storage download failed: "
            f"{response.status_code} "
            f"{response.text}"
        )

    return response.content


# ============================================================
# STORAGE DELETE
# ============================================================

def storage_delete(storage_path):

    if not storage_path:
        return False

    url = (
        f"{SUPABASE_URL}/storage/v1/object/"
        f"{STORAGE_BUCKET}"
    )

    headers = {
        "Authorization": (
            f"Bearer {SUPABASE_SERVICE_KEY}"
        ),
        "apikey": SUPABASE_SERVICE_KEY,
        "Content-Type": "application/json",
    }

    try:

        response = requests.delete(
            url,
            headers=headers,
            json={
                "prefixes": [
                    storage_path
                ]
            },
            timeout=30,
        )

        return response.ok

    except requests.RequestException:

        logging.exception(
            "Storage deletion failed."
        )

        return False


# ============================================================
# SESSION
# ============================================================

def current_user():
    return session.get("user")


def current_email():

    user = current_user()

    if not user:
        return None

    return user.get("email")


def current_user_id():

    user = current_user()

    if not user:
        return None

    return user.get("id")


# ============================================================
# LOGIN REQUIRED
# ============================================================

def login_required(function):

    @wraps(function)
    def wrapper(*args, **kwargs):

        if not current_user():

            flash(
                "Please login first.",
                "warning",
            )

            return redirect(
                url_for("login")
            )

        return function(
            *args,
            **kwargs
        )

    return wrapper


# ============================================================
# ADMIN REQUIRED
# ============================================================

def admin_required(function):

    @wraps(function)
    def wrapper(*args, **kwargs):

        user = current_user()

        if not user:

            flash(
                "Please login first.",
                "warning",
            )

            return redirect(
                url_for("login")
            )

        if not user.get("is_admin"):

            flash(
                "Administrator access required.",
                "danger",
            )

            return redirect(
                url_for("dashboard")
            )

        return function(
            *args,
            **kwargs
        )

    return wrapper


# ============================================================
# USERS TABLE
#
# THIS IS THE IMPORTANT FIX.
#
# We no longer use the broken profiles table.
# ============================================================

def find_user_by_email(email):

    email = (
        email
        or ""
    ).strip().lower()

    if not email:
        return None

    rows = db_select(
        "users",
        filters={
            "email": f"eq.{email}",
        },
        limit=1,
    )

    if rows:
        return rows[0]

    return None


def find_user_by_id(user_id):

    if not user_id:
        return None

    rows = db_select(
        "users",
        filters={
            "id": f"eq.{user_id}",
        },
        limit=1,
    )

    if rows:
        return rows[0]

    return None


# ============================================================
# CREATE USER
# ============================================================

def create_user(
    name,
    email,
    password,
):

    name = (
        name
        or ""
    ).strip()

    email = (
        email
        or ""
    ).strip().lower()

    password = password or ""

    if not name:
        raise ValueError(
            "Full name is required."
        )

    if not email:
        raise ValueError(
            "Email is required."
        )

    if not password:
        raise ValueError(
            "Password is required."
        )

    if len(password) < 6:
        raise ValueError(
            "Password must contain at least 6 characters."
        )

    # Check existing user.
    existing = find_user_by_email(
        email
    )

    if existing:

        raise ValueError(
            "An account with that email already exists."
        )

    user_id = str(
        uuid.uuid4()
    )

    password_hash = generate_password_hash(
        password
    )

    # IMPORTANT:
    #
    # Insert directly into users.
    #
    # We only use columns that definitely exist
    # according to the schema you provided.

    user_data = {
        "id": user_id,
        "name": name,
        "full_name": name,
        "email": email,
        "role": "student",
        "password_hash": password_hash,
        "is_active": True,
        "is_admin": False,
        "created_at": utc_now(),
        "updated_at": utc_now(),
    }

    try:

        result = db_insert(
            "users",
            user_data,
            returning=True,
        )

    except Exception as exc:

        message = str(exc)

        # Duplicate email protection.
        if (
            "duplicate"
            in message.lower()
            or "23505"
            in message
        ):

            raise ValueError(
                "An account with that email already exists."
            )

        raise

    if not result:

        raise RuntimeError(
            "User account was not created."
        )

    if isinstance(result, list):

        return result[0]

    return result


# ============================================================
# VERIFY USER
# ============================================================

def verify_user(
    email,
    password,
):

    user = find_user_by_email(
        email
    )

    if not user:
        return None

    # Disabled users cannot login.
    is_active = user.get(
        "is_active",
        True,
    )

    if is_active is False:
        return None

    password_hash = user.get(
        "password_hash"
    )

    if not password_hash:
        return None

    try:

        valid = check_password_hash(
            password_hash,
            password,
        )

    except Exception:

        logging.exception(
            "Password verification failed."
        )

        return None

    if not valid:
        return None

    return user


# ============================================================
# SESSION USER BUILDER
# ============================================================

def build_session_user(user):

    if not user:
        return None

    role = str(
        user.get(
            "role",
            "student",
        )
    ).lower()

    is_admin = bool(
        user.get(
            "is_admin",
            False,
        )
    )

    if role == "admin":
        is_admin = True

    return {
        "id": user.get("id"),
        "email": user.get("email"),
        "full_name": (
            user.get("full_name")
            or user.get("name")
            or "Student"
        ),
        "role": role,
        "is_admin": is_admin,
    }


# ============================================================
# ACTIVITY LOG
# ============================================================

def log_activity(
    action,
    description="",
):

    user = current_user()

    email = None

    user_id = None

    if user:

        email = user.get(
            "email"
        )

        user_id = user.get(
            "id"
        )

    data = {
        "action": action,
        "description": description,
        "email": email,
        "created_at": utc_now(),
    }

    # user_id is only added if useful.
    # This avoids breaking installations where
    # activity_logs does not contain user_id.

    if user_id:
        data["user_id"] = user_id

    try:

        db_insert(
            "activity_logs",
            data,
            returning=False,
        )

    except Exception:

        # Activity logging must NEVER break
        # registration, login or assignment work.

        logging.warning(
            "Activity log skipped because "
            "activity_logs may not exist."
        )


# ============================================================
# ASSIGNMENT FILE
# ============================================================

def save_assignment_file(
    assignment_id,
    original_filename,
    storage_path,
    content_type,
    file_size,
    file_role="question",
):

    data = {
        "assignment_id": assignment_id,
        "original_filename": original_filename,
        "storage_path": storage_path,
        "content_type": content_type,
        "file_size": file_size,
        "file_role": file_role,
        "created_at": utc_now(),
    }

    result = db_insert(
        "assignment_files",
        data,
        returning=True,
    )

    if isinstance(result, list) and result:
        return result[0]

    return data


# ============================================================
# PDF GENERATOR
# ============================================================

def build_pdf(
    title,
    student_name,
    question,
    answer,
):

    buffer = io.BytesIO()

    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50,
    )

    styles = getSampleStyleSheet()

    title_style = styles["Title"]

    title_style.alignment = TA_CENTER

    heading = styles["Heading2"]

    normal = styles["BodyText"]

    story = []

    story.append(
        Paragraph(
            "KOJA AFRICA",
            title_style,
        )
    )

    story.append(
        Spacer(1, 15)
    )

    story.append(
        Paragraph(
            clean(
                title
                or "Assignment Answer"
            ),
            heading,
        )
    )

    story.append(
        Spacer(1, 10)
    )

    story.append(
        Paragraph(
            "<b>Student:</b> "
            + clean(student_name),
            normal,
        )
    )

    story.append(
        Spacer(1, 15)
    )

    story.append(
        Paragraph(
            "Question",
            heading,
        )
    )

    question_text = clean(
        question
        or ""
    ).replace(
        "\n",
        "<br/>",
    )

    story.append(
        Paragraph(
            question_text
            or "No question supplied.",
            normal,
        )
    )

    story.append(
        Spacer(1, 20)
    )

    story.append(
        Paragraph(
            "Answer",
            heading,
        )
    )

    answer_text = clean(
        answer
        or ""
    ).replace(
        "\n",
        "<br/>",
    )

    story.append(
        Paragraph(
            answer_text
            or "No answer supplied.",
            normal,
        )
    )

    story.append(
        Spacer(1, 30)
    )

    story.append(
        Paragraph(
            "Generated by KOJA AFRICA",
            normal,
        )
    )

    document.build(
        story
    )

    buffer.seek(0)

    return buffer.getvalue()


# ============================================================
# WORD GENERATOR
# ============================================================

def build_docx(
    title,
    student_name,
    question,
    answer,
):

    document = Document()

    document.add_heading(
        "KOJA AFRICA",
        0,
    )

    document.add_heading(
        title
        or "Assignment Answer",
        level=1,
    )

    document.add_paragraph(
        "Student: "
        + (
            student_name
            or ""
        )
    )

    document.add_heading(
        "Question",
        level=2,
    )

    document.add_paragraph(
        question
        or ""
    )

    document.add_heading(
        "Answer",
        level=2,
    )

    document.add_paragraph(
        answer
        or ""
    )

    document.add_paragraph(
        "Generated by KOJA AFRICA"
    )

    buffer = io.BytesIO()

    document.save(
        buffer
    )

    buffer.seek(0)

    return buffer.getvalue()


# ============================================================
# EXCEL GENERATOR
# ============================================================

def build_xlsx(
    title,
    student_name,
    question,
    answer,
):

    workbook = Workbook()

    sheet = workbook.active

    sheet.title = "Assignment Answer"

    sheet["A1"] = "KOJA AFRICA"

    sheet["A2"] = "Assignment"

    sheet["B2"] = (
        title
        or ""
    )

    sheet["A3"] = "Student"

    sheet["B3"] = (
        student_name
        or ""
    )

    sheet["A5"] = "Question"

    sheet["B5"] = (
        question
        or ""
    )

    sheet["A7"] = "Answer"

    sheet["B7"] = (
        answer
        or ""
    )

    sheet.column_dimensions[
        "A"
    ].width = 25

    sheet.column_dimensions[
        "B"
    ].width = 100

    buffer = io.BytesIO()

    workbook.save(
        buffer
    )

    buffer.seek(0)

    return buffer.getvalue()


# ============================================================
# HTML
# ============================================================

BASE_HTML = """
<!doctype html>

<html lang="en">

<head>

<meta charset="utf-8">

<meta name="viewport"
      content="width=device-width, initial-scale=1">

<title>{{ title }} - KOJA AFRICA</title>

<style>

* {
    box-sizing: border-box;
}

body {
    margin: 0;
    font-family:
        Arial,
        Helvetica,
        sans-serif;
    background: #f4f7fb;
    color: #172033;
}

.nav {
    background: #0b3d91;
    color: white;
    padding: 14px 18px;
}

.nav-inner {
    max-width: 1100px;
    margin: auto;
    display: flex;
    justify-content: space-between;
    align-items: center;
    gap: 15px;
    flex-wrap: wrap;
}

.brand {
    font-size: 23px;
    font-weight: bold;
}

.nav-links {
    display: flex;
    flex-wrap: wrap;
    align-items: center;
}

.nav a {
    color: white;
    text-decoration: none;
    margin: 5px 7px;
    display: inline-block;
}

.nav a:hover {
    text-decoration: underline;
}

.container {
    max-width: 1100px;
    margin: 25px auto;
    padding: 0 15px;
}

.card {
    background: white;
    border-radius: 14px;
    padding: 22px;
    margin-bottom: 18px;
    box-shadow:
        0 2px 12px
        rgba(0,0,0,.07);
}

h1,
h2,
h3 {
    margin-top: 0;
}

label {
    display: block;
    font-weight: bold;
    margin-top: 8px;
}

input,
textarea,
select {
    width: 100%;
    padding: 13px;
    margin-top: 6px;
    margin-bottom: 14px;
    border:
        1px solid #ccd3df;
    border-radius: 8px;
    font-size: 16px;
    background: white;
}

textarea {
    min-height: 180px;
    resize: vertical;
}

button,
.btn {
    display: inline-block;
    background: #0b3d91;
    color: white;
    border: 0;
    border-radius: 8px;
    padding: 11px 16px;
    cursor: pointer;
    text-decoration: none;
    margin: 4px;
    font-size: 15px;
}

button:hover,
.btn:hover {
    opacity: .9;
}

.btn-green {
    background: #138a4b;
}

.btn-red {
    background: #b42318;
}

.btn-dark {
    background: #172033;
}

.grid {
    display: grid;
    grid-template-columns:
        repeat(
            auto-fit,
            minmax(210px, 1fr)
        );
    gap: 15px;
}

.stat {
    padding: 20px;
    border-radius: 10px;
    background: #eef4ff;
}

.stat strong {
    display: block;
    font-size: 30px;
    margin-bottom: 5px;
}

table {
    width: 100%;
    border-collapse: collapse;
}

th,
td {
    padding: 10px;
    border-bottom:
        1px solid #ddd;
    text-align: left;
}

.badge {
    display: inline-block;
    padding: 5px 9px;
    border-radius: 15px;
    background: #e8eef8;
}

.flash {
    padding: 13px;
    margin-bottom: 12px;
    border-radius: 8px;
    background: #eef4ff;
}

.footer {
    text-align: center;
    padding: 30px;
    color: #687386;
}

.small {
    color: #687386;
    font-size: 14px;
}

.answer-box {
    background: #f7f9fc;
    border-radius: 10px;
    padding: 18px;
    line-height: 1.7;
}

@media(max-width:600px) {

    .nav-inner {
        align-items: flex-start;
    }

    .nav-links {
        width: 100%;
    }

    .nav a {
        margin: 5px 6px;
    }

    table {
        display: block;
        overflow-x: auto;
    }

    .card {
        padding: 18px;
    }

}

</style>

</head>

<body>

<div class="nav">

<div class="nav-inner">

<div class="brand">
KOJA AFRICA
</div>

<div class="nav-links">

<a href="{{ url_for('home') }}">
Home
</a>

{% if session.get('user') %}

<a href="{{ url_for('dashboard') }}">
Dashboard
</a>

<a href="{{ url_for('new_assignment') }}">
Ask Question
</a>

<a href="{{ url_for('assignments') }}">
My Assignments
</a>

<a href="{{ url_for('notifications') }}">
Notifications
</a>

{% if session.get('user', {}).get('is_admin') %}

<a href="{{ url_for('admin_dashboard') }}">
Admin
</a>

{% endif %}

<a href="{{ url_for('logout') }}">
Logout
</a>

{% else %}

<a href="{{ url_for('login') }}">
Login
</a>

<a href="{{ url_for('register') }}">
Register
</a>

{% endif %}

</div>

</div>

</div>


<div class="container">

{% with messages =
    get_flashed_messages(
        with_categories=true
    )
%}

{% for category, message in messages %}

<div class="flash">
{{ message }}
</div>

{% endfor %}

{% endwith %}

{{ body|safe }}

</div>


<div class="footer">

KOJA AFRICA<br>

Assignment Questions • Academic Answers • Learning Resources

</div>

</body>

</html>
"""


# ============================================================
# PAGE RENDERER
# ============================================================

def page(
    title,
    body,
):
    return render_template_string(
        BASE_HTML,
        title=title,
        body=body,
    )


# ============================================================
# HOME
# ============================================================

@app.route("/")
def home():

    body = """

    <div class="card">

    <h1>KOJA AFRICA</h1>

    <p>
    Assignment Questions • Academic Answers •
    Learning Resources
    </p>

    <p>
    Submit academic questions, upload assignments
    and receive completed answers.
    </p>

    <a class="btn"
       href="/register">
       Create Student Account
    </a>

    <a class="btn btn-dark"
       href="/login">
       Login
    </a>

    </div>


    <div class="grid">

    <div class="card">

    <h3>Ask Questions</h3>

    <p>
    Send an academic question directly
    through your KOJA account.
    </p>

    </div>


    <div class="card">

    <h3>Upload Assignments</h3>

    <p>
    Upload PDF, Word, Excel, images or text files.
    </p>

    </div>


    <div class="card">

    <h3>Download Answers</h3>

    <p>
    Completed answers can be downloaded in
    PDF, Word and Excel formats.
    </p>

    </div>

    </div>

    """

    return page(
        "Home",
        body,
    )


# ============================================================
# REGISTER
# ============================================================

@app.route(
    "/register",
    methods=["GET", "POST"],
)
def register():

    if request.method == "POST":

        name = request.form.get(
            "name",
            "",
        ).strip()

        email = request.form.get(
            "email",
            "",
        ).strip().lower()

        password = request.form.get(
            "password",
            "",
        )

        if not name:

            flash(
                "Full name is required.",
                "warning",
            )

            return redirect(
                url_for("register")
            )

        if not email:

            flash(
                "Email is required.",
                "warning",
            )

            return redirect(
                url_for("register")
            )

        if len(password) < 6:

            flash(
                "Password must contain at least 6 characters.",
                "warning",
            )

            return redirect(
                url_for("register")
            )

        try:

            user = create_user(
                name=name,
                email=email,
                password=password,
            )

            session.clear()

            session["user"] = (
                build_session_user(
                    user
                )
            )

            log_activity(
                "register",
                "New student account created.",
            )

            flash(
                "Account created successfully.",
                "success",
            )

            return redirect(
                url_for("dashboard")
            )

        except ValueError as exc:

            flash(
                str(exc),
                "warning",
            )

        except Exception as exc:

            logging.exception(
                "Registration error."
            )

            flash(
                f"Registration error: {exc}",
                "danger",
            )

    body = """

    <div class="card">

    <h1>Create Account</h1>

    <form method="post">

    <label>Full name</label>

    <input
        name="name"
        required
        autocomplete="name"
        placeholder="Your full name"
    >

    <label>Email</label>

    <input
        type="email"
        name="email"
        required
        autocomplete="email"
        placeholder="you@example.com"
    >

    <label>Password</label>

    <input
        type="password"
        name="password"
        required
        minlength="6"
        autocomplete="new-password"
        placeholder="At least 6 characters"
    >

    <button type="submit">
    Create Account
    </button>

    </form>

    <p>
    Already have an account?
    <a href="/login">Login</a>
    </p>

    </div>

    """

    return page(
        "Register",
        body,
    )


# ============================================================
# LOGIN
# ============================================================

@app.route(
    "/login",
    methods=["GET", "POST"],
)
def login():

    if request.method == "POST":

        email = request.form.get(
            "email",
            "",
        ).strip().lower()

        password = request.form.get(
            "password",
            "",
        )

        if not email or not password:

            flash(
                "Enter email and password.",
                "warning",
            )

            return redirect(
                url_for("login")
            )

        try:

            # ==================================================
            # ADMIN ENVIRONMENT LOGIN
            # ==================================================

            if (
                email == ADMIN_EMAIL
                and password == ADMIN_PASSWORD
            ):

                session.clear()

                session["user"] = {
                    "id": "admin",
                    "email": ADMIN_EMAIL,
                    "full_name": "KOJA Administrator",
                    "role": "admin",
                    "is_admin": True,
                }

                log_activity(
                    "admin_login",
                    "Administrator logged in.",
                )

                return redirect(
                    url_for(
                        "admin_dashboard"
                    )
                )

            # ==================================================
            # NORMAL USER
            # ==================================================

            user = verify_user(
                email,
                password,
            )

            if not user:

                flash(
                    "Invalid email or password.",
                    "danger",
                )

                return redirect(
                    url_for("login")
                )

            session.clear()

            session["user"] = (
                build_session_user(
                    user
                )
            )

            log_activity(
                "login",
                "User logged in.",
            )

            if session["user"].get(
                "is_admin"
            ):

                return redirect(
                    url_for(
                        "admin_dashboard"
                    )
                )

            return redirect(
                url_for("dashboard")
            )

        except Exception as exc:

            logging.exception(
                "Login error."
            )

            flash(
                f"Login error: {exc}",
                "danger",
            )

    body = """

    <div class="card">

    <h1>Login</h1>

    <form method="post">

    <label>Email</label>

    <input
        type="email"
        name="email"
        required
        autocomplete="email"
        placeholder="you@example.com"
    >

    <label>Password</label>

    <input
        type="password"
        name="password"
        required
        autocomplete="current-password"
        placeholder="Your password"
    >

    <button type="submit">
    Login
    </button>

    </form>

    <p>
    Don't have an account?
    <a href="/register">
    Create one
    </a>
    </p>

    </div>

    """

    return page(
        "Login",
        body,
    )


# ============================================================
# LOGOUT
# ============================================================

@app.route("/logout")
def logout():

    session.clear()

    flash(
        "You have been logged out.",
        "success",
    )

    return redirect(
        url_for("home")
    )


# ============================================================
# DASHBOARD
# ============================================================

@app.route("/dashboard")
@login_required
def dashboard():

    email = current_email()

    rows = []

    try:

        rows = db_select(
            "assignments",
            filters={
                "student_email":
                    f"eq.{email}"
            },
            limit=100,
            order="created_at.desc",
        )

    except Exception as exc:

        logging.exception(
            "Dashboard assignment query failed."
        )

        flash(
            f"Could not load assignments: {exc}",
            "danger",
        )

    pending = 0
    processing = 0
    completed = 0

    for item in rows:

        status = str(
            item.get(
                "status",
                "",
            )
        ).lower()

        if status == "pending":
            pending += 1

        elif status == "processing":
            processing += 1

        elif status == "completed":
            completed += 1

    username = clean(
        current_user().get(
            "full_name",
            "Student",
        )
    )

    body = f"""

    <div class="card">

    <h1>
    Welcome, {username}
    </h1>

    <p>
    Submit questions and manage your assignments.
    </p>

    </div>


    <div class="grid">

    <div class="stat">
    <strong>{len(rows)}</strong>
    Total Assignments
    </div>

    <div class="stat">
    <strong>{pending}</strong>
    Pending
    </div>

    <div class="stat">
    <strong>{processing}</strong>
    Processing
    </div>

    <div class="stat">
    <strong>{completed}</strong>
    Completed
    </div>

    </div>


    <div class="card">

    <h2>Quick Actions</h2>

    <a class="btn"
       href="/assignment/new">
       Ask Question / Upload Assignment
    </a>

    <a class="btn btn-dark"
       href="/assignments">
       View My Assignments
    </a>

    <a class="btn"
       href="/notifications">
       Notifications
    </a>

    </div>

    """

    return page(
        "Dashboard",
        body,
    )


# ============================================================
# NEW ASSIGNMENT
# ============================================================

@app.route(
    "/assignment/new",
    methods=["GET", "POST"],
)
@login_required
def new_assignment():

    if request.method == "POST":

        title = request.form.get(
            "title",
            "",
        ).strip()

        subject = request.form.get(
            "subject",
            "",
        ).strip()

        question = request.form.get(
            "question",
            "",
        ).strip()

        uploaded_file = request.files.get(
            "question_file"
        )

        if not title:
            title = "Assignment Question"

        has_file = bool(
            uploaded_file
            and uploaded_file.filename
        )

        if not question and not has_file:

            flash(
                "Enter a question or upload a question file.",
                "warning",
            )

            return redirect(
                url_for(
                    "new_assignment"
                )
            )

        # Validate file BEFORE creating database
        # record so bad files do not create
        # empty assignments.

        file_bytes = None
        original_name = None
        extension = None
        content_type = None

        if has_file:

            if not allowed_file(
                uploaded_file.filename
            ):

                flash(
                    "Unsupported file type.",
                    "danger",
                )

                return redirect(
                    url_for(
                        "new_assignment"
                    )
                )

            file_bytes = uploaded_file.read()

            if len(file_bytes) > MAX_FILE_SIZE:

                flash(
                    "File is too large. Maximum is 15 MB.",
                    "danger",
                )

                return redirect(
                    url_for(
                        "new_assignment"
                    )
                )

            original_name = secure_filename(
                uploaded_file.filename
            )

            if (
                not original_name
                or "." not in original_name
            ):

                flash(
                    "Invalid file name.",
                    "danger",
                )

                return redirect(
                    url_for(
                        "new_assignment"
                    )
                )

            extension = (
                original_name
                .rsplit(".", 1)[-1]
                .lower()
            )

            content_type = (
                uploaded_file.content_type
                or "application/octet-stream"
            )

        try:

            assignment_data = {
                "student_email":
                    current_email(),

                "title":
                    title,

                "subject":
                    subject,

                "question":
                    question,

                "status":
                    "Pending",

                "created_at":
                    utc_now(),

                "updated_at":
                    utc_now(),
            }

            try:

                result = db_insert(
                    "assignments",
                    assignment_data,
                    returning=True,
                )

            except Exception:

                logging.exception(
                    "Full assignment insert failed."
                )

                # Fallback for databases where
                # subject or timestamps may not exist.

                fallback = {
                    "student_email":
                        current_email(),

                    "title":
                        title,

                    "question":
                        question,

                    "status":
                        "Pending",
                }

                result = db_insert(
                    "assignments",
                    fallback,
                    returning=True,
                )

            if not result:

                raise RuntimeError(
                    "Assignment was not created."
                )

            assignment = (
                result[0]
                if isinstance(
                    result,
                    list
                )
                else result
            )

            assignment_id = assignment.get(
                "id"
            )

            if not assignment_id:

                raise RuntimeError(
                    "Assignment was created but no ID was returned."
                )

            # ==================================================
            # UPLOAD QUESTION FILE
            # ==================================================

            if has_file:

                storage_path = (
                    f"questions/"
                    f"{assignment_id}/"
                    f"{uuid.uuid4().hex}."
                    f"{extension}"
                )

                storage_upload(
                    file_bytes,
                    storage_path,
                    content_type,
                )

                try:

                    save_assignment_file(
                        assignment_id,
                        original_name,
                        storage_path,
                        content_type,
                        len(file_bytes),
                        "question",
                    )

                except Exception:

                    # If database file record fails,
                    # remove the storage object.
                    storage_delete(
                        storage_path
                    )

                    raise

            log_activity(
                "assignment_created",
                (
                    f"Assignment "
                    f"{assignment_id} submitted."
                ),
            )

            flash(
                "Assignment submitted successfully.",
                "success",
            )

            return redirect(
                url_for(
                    "assignment_detail",
                    assignment_id=assignment_id,
                )
            )

        except Exception as exc:

            logging.exception(
                "Assignment creation failed."
            )

            flash(
                f"Could not submit assignment: {exc}",
                "danger",
            )

    body = """

    <div class="card">

    <h1>Assignment Request</h1>

    <p>
    Ask a question or upload your assignment.
    </p>

    <form
        method="post"
        enctype="multipart/form-data"
    >

    <label>Assignment title</label>

    <input
        name="title"
        placeholder="Example: Chemistry Assignment 1"
    >

    <label>Subject</label>

    <input
        name="subject"
        placeholder="Example: Chemistry"
    >

    <label>Ask your question</label>

    <textarea
        name="question"
        placeholder="Type your assignment question here..."
    ></textarea>

    <label>
    Upload question
    </label>

    <input
        type="file"
        name="question_file"
        accept=".pdf,.doc,.docx,.xls,.xlsx,.png,.jpg,.jpeg,.webp,.txt"
    >

    <p class="small">
    Supported: PDF, Word, Excel, images and text.
    Maximum: 15 MB.
    </p>

    <button type="submit">
    Submit Assignment
    </button>

    </form>

    </div>

    """

    return page(
        "Assignment Request",
        body,
    )


# ============================================================
# ASSIGNMENTS
# ============================================================

@app.route("/assignments")
@login_required
def assignments():

    rows = []

    try:

        rows = db_select(
            "assignments",
            filters={
                "student_email":
                    f"eq.{current_email()}"
            },
            limit=200,
            order="created_at.desc",
        )

    except Exception as exc:

        logging.exception(
            "Assignment listing failed."
        )

        flash(
            f"Could not load assignments: {exc}",
            "danger",
        )

    html = """

    <div class="card">

    <h1>My Assignments</h1>

    <a class="btn"
       href="/assignment/new">
       New Assignment
    </a>

    </div>

    """

    if not rows:

        html += """

        <div class="card">

        <p>
        No assignments submitted yet.
        </p>

        </div>

        """

    else:

        html += """

        <div class="card">

        <table>

        <tr>

        <th>Title</th>

        <th>Subject</th>

        <th>Status</th>

        <th>Date</th>

        <th>Action</th>

        </tr>

        """

        for item in rows:

            assignment_id = clean(
                item.get("id")
            )

            title = clean(
                item.get(
                    "title",
                    "Assignment",
                )
            )

            subject = clean(
                item.get(
                    "subject",
                    "",
                )
            )

            status = clean(
                item.get(
                    "status",
                    "Pending",
                )
            )

            created = clean(
                item.get(
                    "created_at",
                    "",
                )
            )

            html += f"""

            <tr>

            <td>{title}</td>

            <td>{subject}</td>

            <td>
            <span class="badge">
            {status}
            </span>
            </td>

            <td>{created}</td>

            <td>

            <a class="btn"
               href="/assignment/{assignment_id}">
               Open
            </a>

            </td>

            </tr>

            """

        html += """

        </table>

        </div>

        """

    return page(
        "My Assignments",
        html,
    )


# ============================================================
# ASSIGNMENT DETAIL
# ============================================================

@app.route(
    "/assignment/<assignment_id>"
)
@login_required
def assignment_detail(
    assignment_id
):

    try:

        rows = db_select(
            "assignments",
            filters={
                "id":
                    f"eq.{assignment_id}",
                "student_email":
                    f"eq.{current_email()}",
            },
            limit=1,
        )

        if not rows:
            abort(404)

        assignment = rows[0]

        try:

            files = db_select(
                "assignment_files",
                filters={
                    "assignment_id":
                        f"eq.{assignment_id}"
                },
                limit=100,
                order="created_at.desc",
            )

        except Exception:

            logging.warning(
                "assignment_files table unavailable."
            )

            files = []

    except Exception as exc:

        logging.exception(
            "Assignment detail error."
        )

        flash(
            f"Could not load assignment: {exc}",
            "danger",
        )

        return redirect(
            url_for("assignments")
        )

    title = clean(
        assignment.get(
            "title",
            "Assignment",
        )
    )

    subject = clean(
        assignment.get(
            "subject",
            "",
        )
    )

    status = clean(
        assignment.get(
            "status",
            "Pending",
        )
    )

    question = clean(
        assignment.get(
            "question",
            "",
        )
    ).replace(
        "\n",
        "<br>",
    )

    html = f"""

    <div class="card">

    <h1>{title}</h1>

    <p>
    <b>Subject:</b>
    {subject}
    </p>

    <p>
    <b>Status:</b>

    <span class="badge">
    {status}
    </span>

    </p>

    <h2>Question</h2>

    <div class="answer-box">
    {question or "No question text supplied."}
    </div>

    </div>

    """

    if files:

        html += """

        <div class="card">

        <h2>Files</h2>

        """

        for file_record in files:

            file_id = clean(
                file_record.get(
                    "id"
                )
            )

            name = clean(
                file_record.get(
                    "original_filename",
                    "File",
                )
            )

            role = clean(
                file_record.get(
                    "file_role",
                    "question",
                )
            )

            html += f"""

            <p>

            <b>{name}</b>
            ({role})

            <a class="btn"
               href="/file/{file_id}">
               Download
            </a>

            </p>

            """

        html += """

        </div>

        """

    status_lower = str(
        assignment.get(
            "status",
            "",
        )
    ).lower()

    if status_lower == "completed":

        answer = clean(
            assignment.get(
                "answer",
                "",
            )
        ).replace(
            "\n",
            "<br>",
        )

        html += f"""

        <div class="card">

        <h2>Completed Answer</h2>

        <div class="answer-box">

        {answer}

        </div>

        <br>

        <a class="btn"
           href="/answer/{assignment_id}/pdf">
           Download PDF
        </a>

        <a class="btn btn-green"
           href="/answer/{assignment_id}/docx">
           Download Word
        </a>

        <a class="btn btn-dark"
           href="/answer/{assignment_id}/xlsx">
           Download Excel
        </a>

        </div>

        """

    return page(
        "Assignment",
        html,
    )


# ============================================================
# QUESTION FILE DOWNLOAD
# ============================================================

@app.route(
    "/file/<file_id>"
)
@login_required
def download_question_file(
    file_id
):

    try:

        files = db_select(
            "assignment_files",
            filters={
                "id":
                    f"eq.{file_id}"
            },
            limit=1,
        )

        if not files:
            abort(404)

        file_record = files[0]

        assignment_id = file_record.get(
            "assignment_id"
        )

        if not assignment_id:
            abort(404)

        assignment_rows = db_select(
            "assignments",
            filters={
                "id":
                    f"eq.{assignment_id}"
            },
            limit=1,
        )

        if not assignment_rows:
            abort(404)

        assignment = assignment_rows[0]

        user = current_user()

        is_admin = bool(
            user.get(
                "is_admin",
                False,
            )
        )

        if (
            not is_admin
            and assignment.get(
                "student_email"
            ) != current_email()
        ):
            abort(403)

        storage_path = file_record.get(
            "storage_path"
        )

        content = storage_download(
            storage_path
        )

        filename = (
            file_record.get(
                "original_filename"
            )
            or "download"
        )

        content_type = (
            file_record.get(
                "content_type"
            )
            or "application/octet-stream"
        )

        return send_file(
            io.BytesIO(content),
            as_attachment=True,
            download_name=filename,
            mimetype=content_type,
        )

    except Exception as exc:

        logging.exception(
            "File download error."
        )

        flash(
            f"File download failed: {exc}",
            "danger",
        )

        return redirect(
            url_for("assignments")
        )


# ============================================================
# ADMIN DASHBOARD
# ============================================================

@app.route("/admin")
@admin_required
def admin_dashboard():

    assignments_rows = []

    users = []

    try:

        assignments_rows = db_select(
            "assignments",
            limit=500,
            order="created_at.desc",
        )

    except Exception as exc:

        logging.exception(
            "Admin assignment query failed."
        )

        flash(
            f"Could not load assignments: {exc}",
            "danger",
        )

    try:

        users = db_select(
            "users",
            limit=500,
            order="created_at.desc",
        )

    except Exception as exc:

        logging.exception(
            "Admin user query failed."
        )

        flash(
            f"Could not load users: {exc}",
            "danger",
        )

    pending = 0
    processing = 0
    completed = 0

    for item in assignments_rows:

        status = str(
            item.get(
                "status",
                "",
            )
        ).lower()

        if status == "pending":
            pending += 1

        elif status == "processing":
            processing += 1

        elif status == "completed":
            completed += 1

    html = f"""

    <div class="card">

    <h1>KOJA AFRICA Admin</h1>

    <p>
    Administrator:
    {clean(current_email())}
    </p>

    <div class="grid">

    <div class="stat">
    <strong>{len(users)}</strong>
    Users
    </div>

    <div class="stat">
    <strong>{len(assignments_rows)}</strong>
    Assignments
    </div>

    <div class="stat">
    <strong>{pending}</strong>
    Pending
    </div>

    <div class="stat">
    <strong>{processing}</strong>
    Processing
    </div>

    <div class="stat">
    <strong>{completed}</strong>
    Completed
    </div>

    </div>

    </div>


    <div class="card">

    <h2>Assignment Requests</h2>

    """

    if not assignments_rows:

        html += """

        <p>
        No assignments.
        </p>

        """

    else:

        html += """

        <table>

        <tr>

        <th>Student</th>

        <th>Title</th>

        <th>Subject</th>

        <th>Status</th>

        <th>Action</th>

        </tr>

        """

        for item in assignments_rows:

            aid = clean(
                item.get(
                    "id",
                    "",
                )
            )

            student = clean(
                item.get(
                    "student_email",
                    "",
                )
            )

            title = clean(
                item.get(
                    "title",
                    "",
                )
            )

            subject = clean(
                item.get(
                    "subject",
                    "",
                )
            )

            status = clean(
                item.get(
                    "status",
                    "Pending",
                )
            )

            html += f"""

            <tr>

            <td>
            {student}
            </td>

            <td>
            {title}
            </td>

            <td>
            {subject}
            </td>

            <td>

            <span class="badge">
            {status}
            </span>

            </td>

            <td>

            <a class="btn"
               href="/admin/assignment/{aid}">
               Process
            </a>

            </td>

            </tr>

            """

        html += """

        </table>

        """

    html += """

    </div>


    <div class="card">

    <h2>Registered Users</h2>

    """

    if users:

        html += """

        <table>

        <tr>

        <th>Name</th>

        <th>Email</th>

        <th>Role</th>

        <th>Active</th>

        </tr>

        """

        for user in users:

            name = clean(
                user.get(
                    "full_name"
                )
                or user.get(
                    "name"
                )
                or ""
            )

            email = clean(
                user.get(
                    "email",
                    "",
                )
            )

            role = clean(
                user.get(
                    "role",
                    "student",
                )
            )

            active = (
                "Yes"
                if user.get(
                    "is_active",
                    True,
                )
                else "No"
            )

            html += f"""

            <tr>

            <td>{name}</td>

            <td>{email}</td>

            <td>{role}</td>

            <td>{active}</td>

            </tr>

            """

        html += """

        </table>

        """

    else:

        html += """

        <p>
        No registered users found.
        </p>

        """

    html += """

    </div>

    """

    return page(
        "Admin Dashboard",
        html,
    )


# ============================================================
# ADMIN PROCESS ASSIGNMENT
# ============================================================

@app.route(
    "/admin/assignment/<assignment_id>",
    methods=["GET", "POST"],
)
@admin_required
def admin_assignment(
    assignment_id
):

    try:

        rows = db_select(
            "assignments",
            filters={
                "id":
                    f"eq.{assignment_id}"
            },
            limit=1,
        )

        if not rows:
            abort(404)

        assignment = rows[0]

    except Exception as exc:

        logging.exception(
            "Could not load admin assignment."
        )

        flash(
            f"Could not load assignment: {exc}",
            "danger",
        )

        return redirect(
            url_for("admin_dashboard")
        )

    if request.method == "POST":

        action = request.form.get(
            "action",
            "",
        )

        # ==================================================
        # PROCESSING
        # ==================================================

        if action == "processing":

            try:

                db_update(
                    "assignments",
                    {
                        "id":
                            f"eq.{assignment_id}"
                    },
                    {
                        "status":
                            "Processing",

                        "updated_at":
                            utc_now(),
                    },
                )

                log_activity(
                    "assignment_processing",
                    (
                        f"{assignment_id} "
                        f"moved to processing."
                    ),
                )

                flash(
                    "Assignment marked as processing.",
                    "success",
                )

            except Exception as exc:

                logging.exception(
                    "Could not update processing status."
                )

                flash(
                    f"Could not update status: {exc}",
                    "danger",
                )

            return redirect(
                url_for(
                    "admin_assignment",
                    assignment_id=assignment_id,
                )
            )

        # ==================================================
        # COMPLETE
        # ==================================================

        if action == "complete":

            answer = request.form.get(
                "answer",
                "",
            ).strip()

            if not answer:

                flash(
                    "Enter the answer before completing the assignment.",
                    "warning",
                )

                return redirect(
                    url_for(
                        "admin_assignment",
                        assignment_id=assignment_id,
                    )
                )

            try:

                db_update(
                    "assignments",
                    {
                        "id":
                            f"eq.{assignment_id}"
                    },
                    {
                        "answer":
                            answer,

                        "status":
                            "Completed",

                        "updated_at":
                            utc_now(),

                        "completed_at":
                            utc_now(),
                    },
                )

                # ==================================================
                # NOTIFICATION
                # ==================================================

                notification = {
                    "student_email":
                        assignment.get(
                            "student_email"
                        ),

                    "title":
                        "Assignment Completed",

                    "message":
                        "Your assignment answer is ready for download.",

                    "created_at":
                        utc_now(),
                }

                try:

                    db_insert(
                        "notifications",
                        notification,
                        returning=False,
                    )

                except Exception:

                    logging.warning(
                        "Notification table unavailable."
                    )

                log_activity(
                    "assignment_completed",
                    (
                        f"{assignment_id} "
                        f"completed."
                    ),
                )

                flash(
                    "Assignment completed successfully.",
                    "success",
                )

                return redirect(
                    url_for(
                        "admin_assignment",
                        assignment_id=assignment_id,
                    )
                )

            except Exception as exc:

                logging.exception(
                    "Assignment completion failed."
                )

                flash(
                    f"Could not complete assignment: {exc}",
                    "danger",
                )

    # ==================================================
    # LOAD FILES
    # ==================================================

    try:

        files = db_select(
            "assignment_files",
            filters={
                "assignment_id":
                    f"eq.{assignment_id}"
            },
            limit=100,
            order="created_at.desc",
        )

    except Exception:

        logging.warning(
            "Could not load assignment files."
        )

        files = []

    title = clean(
        assignment.get(
            "title",
            "",
        )
    )

    student_email = clean(
        assignment.get(
            "student_email",
            "",
        )
    )

    subject = clean(
        assignment.get(
            "subject",
            "",
        )
    )

    status = clean(
        assignment.get(
            "status",
            "Pending",
        )
    )

    question = clean(
        assignment.get(
            "question",
            "",
        )
    ).replace(
        "\n",
        "<br>",
    )

    existing_answer = clean(
        assignment.get(
            "answer",
            "",
        )
    )

    html = f"""

    <div class="card">

    <h1>Process Assignment</h1>

    <p>
    <b>Student:</b>
    {student_email}
    </p>

    <p>
    <b>Title:</b>
    {title}
    </p>

    <p>
    <b>Subject:</b>
    {subject}
    </p>

    <p>
    <b>Status:</b>

    <span class="badge">
    {status}
    </span>

    </p>

    <h2>Question</h2>

    <div class="answer-box">

    {question or "No question text supplied."}

    </div>

    """

    if files:

        html += """

        <h2>Uploaded Files</h2>

        """

        for file_record in files:

            file_id = clean(
                file_record.get(
                    "id"
                )
            )

            filename = clean(
                file_record.get(
                    "original_filename",
                    "File",
                )
            )

            html += f"""

            <p>

            <b>{filename}</b>

            <a class="btn"
               href="/file/{file_id}">
               Download
            </a>

            </p>

            """

    html += f"""

    <hr>

    <form method="post">

    <input
        type="hidden"
        name="action"
        value="processing"
    >

    <button type="submit">
    Mark Processing
    </button>

    </form>


    <h2>Write Answer</h2>

    <form method="post">

    <input
        type="hidden"
        name="action"
        value="complete"
    >

    <textarea
        name="answer"
        placeholder="Write the completed academic answer here..."
        required
    >{existing_answer}</textarea>

    <button
        type="submit"
        class="btn-green"
    >
    Complete Assignment
    </button>

    </form>

    </div>

    """

    return page(
        "Process Assignment",
        html,
    )


# ============================================================
# GET COMPLETED ASSIGNMENT
# ============================================================

def get_completed_assignment(
    assignment_id
):

    rows = db_select(
        "assignments",
        filters={
            "id":
                f"eq.{assignment_id}"
        },
        limit=1,
    )

    if not rows:
        abort(404)

    assignment = rows[0]

    user = current_user()

    if not user:
        abort(403)

    is_admin = bool(
        user.get(
            "is_admin",
            False,
        )
    )

    if (
        not is_admin
        and assignment.get(
            "student_email"
        ) != current_email()
    ):

        abort(403)

    if str(
        assignment.get(
            "status",
            "",
        )
    ).lower() != "completed":

        flash(
            "This assignment has not been completed yet.",
            "warning",
        )

        abort(404)

    return assignment


# ============================================================
# DOWNLOAD LOG
# ============================================================

def log_download(
    assignment_id,
    file_type,
):

    data = {
        "assignment_id":
            assignment_id,

        "student_email":
            current_email(),

        "file_type":
            file_type,

        "downloaded_at":
            utc_now(),
    }

    try:

        db_insert(
            "downloads",
            data,
            returning=False,
        )

    except Exception:

        logging.warning(
            "Download logging skipped."
        )


# ============================================================
# ANSWER PDF
# ============================================================

@app.route(
    "/answer/<assignment_id>/pdf"
)
@login_required
def answer_pdf(
    assignment_id
):

    assignment = (
        get_completed_assignment(
            assignment_id
        )
    )

    data = build_pdf(
        assignment.get(
            "title",
            "Assignment Answer",
        ),

        assignment.get(
            "student_email",
            "",
        ),

        assignment.get(
            "question",
            "",
        ),

        assignment.get(
            "answer",
            "",
        ),
    )

    log_download(
        assignment_id,
        "pdf",
    )

    return send_file(
        io.BytesIO(data),
        as_attachment=True,
        download_name=(
            f"KOJA-Answer-{assignment_id}.pdf"
        ),
        mimetype="application/pdf",
    )


# ============================================================
# ANSWER WORD
# ============================================================

@app.route(
    "/answer/<assignment_id>/docx"
)
@login_required
def answer_docx(
    assignment_id
):

    assignment = (
        get_completed_assignment(
            assignment_id
        )
    )

    data = build_docx(
        assignment.get(
            "title",
            "Assignment Answer",
        ),

        assignment.get(
            "student_email",
            "",
        ),

        assignment.get(
            "question",
            "",
        ),

        assignment.get(
            "answer",
            "",
        ),
    )

    log_download(
        assignment_id,
        "docx",
    )

    return send_file(
        io.BytesIO(data),
        as_attachment=True,
        download_name=(
            f"KOJA-Answer-{assignment_id}.docx"
        ),
        mimetype=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )


# ============================================================
# ANSWER EXCEL
# ============================================================

@app.route(
    "/answer/<assignment_id>/xlsx"
)
@login_required
def answer_xlsx(
    assignment_id
):

    assignment = (
        get_completed_assignment(
            assignment_id
        )
    )

    data = build_xlsx(
        assignment.get(
            "title",
            "Assignment Answer",
        ),

        assignment.get(
            "student_email",
            "",
        ),

        assignment.get(
            "question",
            "",
        ),

        assignment.get(
            "answer",
            "",
        ),
    )

    log_download(
        assignment_id,
        "xlsx",
    )

    return send_file(
        io.BytesIO(data),
        as_attachment=True,
        download_name=(
            f"KOJA-Answer-{assignment_id}.xlsx"
        ),
        mimetype=(
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
    )


# ============================================================
# NOTIFICATIONS
# ============================================================

@app.route("/notifications")
@login_required
def notifications():

    rows = []

    try:

        rows = db_select(
            "notifications",
            filters={
                "student_email":
                    f"eq.{current_email()}"
            },
            limit=100,
            order="created_at.desc",
        )

    except Exception as exc:

        logging.exception(
            "Notifications query failed."
        )

        flash(
            f"Could not load notifications: {exc}",
            "danger",
        )

    html = """

    <div class="card">

    <h1>Notifications</h1>

    """

    if not rows:

        html += """

        <p>
        No notifications.
        </p>

        """

    else:

        for notification in rows:

            title = clean(
                notification.get(
                    "title",
                    "Notification",
                )
            )

            message = clean(
                notification.get(
                    "message",
                    "",
                )
            )

            created_at = clean(
                notification.get(
                    "created_at",
                    "",
                )
            )

            html += f"""

            <div class="card">

            <h3>
            {title}
            </h3>

            <p>
            {message}
            </p>

            <small>
            {created_at}
            </small>

            </div>

            """

    html += """

    </div>

    """

    return page(
        "Notifications",
        html,
    )


# ============================================================
# HEALTH CHECK
# ============================================================

@app.route("/health")
def health():

    return jsonify(
        {
            "status": "ok",
            "application": APP_NAME,
            "time": utc_now(),
            "supabase_configured": bool(
                SUPABASE_URL
                and SUPABASE_SERVICE_KEY
            ),
            "storage_bucket": STORAGE_BUCKET,
        }
    )


# ============================================================
# DATABASE HEALTH CHECK
# ============================================================

@app.route("/health/database")
def database_health():

    result = {
        "status": "unknown",
        "supabase": False,
        "users_table": False,
        "message": "",
    }

    try:

        if not SUPABASE_URL:
            raise RuntimeError(
                "SUPABASE_URL is missing."
            )

        if not SUPABASE_SERVICE_KEY:
            raise RuntimeError(
                "SUPABASE_SERVICE_KEY is missing."
            )

        result["supabase"] = True

        users = db_select(
            "users",
            columns="id",
            limit=1,
        )

        result["users_table"] = True

        result["status"] = "ok"

        result["message"] = (
            "Supabase connection and users table are working."
        )

        result["user_rows_checked"] = len(
            users
        )

        return jsonify(result)

    except Exception as exc:

        result["status"] = "error"

        result["message"] = str(exc)

        return jsonify(result), 500


# ============================================================
# ERROR HANDLER 403
# ============================================================

@app.errorhandler(403)
def error_403(error):

    return page(
        "Access Denied",
        """

        <div class="card">

        <h1>403</h1>

        <p>
        You do not have permission to access this page.
        </p>

        <a class="btn"
           href="/">
           Return Home
        </a>

        </div>

        """,
    ), 403


# ============================================================
# ERROR HANDLER 404
# ============================================================

@app.errorhandler(404)
def error_404(error):

    return page(
        "Page Not Found",
        """

        <div class="card">

        <h1>404</h1>

        <p>
        The requested page was not found.
        </p>

        <a class="btn"
           href="/">
           Return Home
        </a>

        </div>

        """,
    ), 404


# ============================================================
# ERROR HANDLER 413
# ============================================================

@app.errorhandler(413)
def error_413(error):

    return page(
        "File Too Large",
        """

        <div class="card">

        <h1>File Too Large</h1>

        <p>
        The maximum upload size is 15 MB.
        </p>

        <a class="btn"
           href="/assignment/new">
           Try Again
        </a>

        </div>

        """,
    ), 413


# ============================================================
# ERROR HANDLER 500
# ============================================================

@app.errorhandler(500)
def error_500(error):

    logging.exception(
        "Unhandled application error."
    )

    return page(
        "Application Error",
        """

        <div class="card">

        <h1>Application Error</h1>

        <p>
        KOJA AFRICA encountered an internal error.
        </p>

        <p>
        Please try again.
        </p>

        <a class="btn"
           href="/">
           Return Home
        </a>

        </div>

        """,
    ), 500


# ============================================================
# START APPLICATION
# ============================================================

if __name__ == "__main__":

    port = int(
        os.environ.get(
            "PORT",
            "10000",
        )
    )

    app.run(
        host="0.0.0.0",
        port=port,
        debug=False,
    )
